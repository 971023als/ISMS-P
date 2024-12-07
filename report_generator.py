import json
from docx import Document
import sys

def generate_report(input_file, output_file):
    try:
        print(f"Generating report from: {input_file}")

        # JSON 데이터 읽기
        with open(input_file, "r") as f:
            data = json.load(f)

        # Word 문서 생성
        doc = Document()
        doc.add_heading("ISMS-P Checklist Analysis Report", level=1)

        for item in data:
            doc.add_heading("Summary:", level=2)
            doc.add_paragraph(f"Total Items: {item['total_items']}")
            doc.add_paragraph(f"Completed: {item['completed']}")
            doc.add_paragraph(f"Pending: {item['pending']}")

        doc.save(output_file)
        print(f"Report saved to: {output_file}")

    except Exception as e:
        print(f"Error generating report: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python report_generator.py --input <input_file> --output <output_file>")
    else:
        input_file = sys.argv[1].split("--input ")[-1]
        output_file = sys.argv[2].split("--output ")[-1]
        generate_report(input_file, output_file)
